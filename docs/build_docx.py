"""기획서_v2.md → 기획서_v2.docx (python-docx).

지원: # ## ### 제목, | | 표, ![](img) 이미지, ``` 코드블록, - 목록,
      **굵게**, `코드`, 인라인.
"""
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = Path(__file__).resolve().parent
lines = (HERE / "기획서_v2.md").read_text("utf-8").splitlines()

ACCENT = RGBColor(0xB8, 0x38, 0x1F)
FONT = "맑은 고딕"

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(10.5)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def set_kr(run):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def add_runs(p, text):
    for part in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); r.bold = True; r.font.color.rgb = RGBColor(0xA8, 0x34, 0x1C)
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1]); r.font.name = "Consolas"
        else:
            r = p.add_run(part)
        set_kr(r)


def cells(row):
    return [c.strip() for c in row.strip().strip("|").split("|")]


i = 0
n = len(lines)
while i < n:
    ln = lines[i]
    s = ln.strip()

    # 코드블록
    if s.startswith("```"):
        i += 1
        code = []
        while i < n and not lines[i].strip().startswith("```"):
            code.append(lines[i]); i += 1
        i += 1
        p = doc.add_paragraph()
        r = p.add_run("\n".join(code)); r.font.name = "Consolas"; r.font.size = Pt(9)
        p.paragraph_format.left_indent = Inches(0.15)
        continue

    # 이미지
    m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", s)
    if m:
        src = HERE / m.group(2)
        if src.exists():
            doc.add_picture(str(src), width=Inches(6.1))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph()
            cr = cap.add_run(m.group(1)); cr.italic = True; cr.font.size = Pt(9)
            cr.font.color.rgb = RGBColor(0x88, 0x88, 0x88); set_kr(cr)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1
        continue

    # 표
    if s.startswith("|") and i + 1 < n and set(lines[i + 1].strip()) <= set("|-: "):
        header = cells(s)
        i += 2
        rows = []
        while i < n and lines[i].strip().startswith("|"):
            rows.append(cells(lines[i])); i += 1
        t = doc.add_table(rows=1, cols=len(header)); t.style = "Light Grid Accent 2"
        for j, h in enumerate(header):
            add_runs(t.rows[0].cells[j].paragraphs[0], h)
            for rn in t.rows[0].cells[j].paragraphs[0].runs:
                rn.bold = True
        for row in rows:
            rc = t.add_row().cells
            for j in range(len(header)):
                add_runs(rc[j].paragraphs[0], row[j] if j < len(row) else "")
        doc.add_paragraph()
        continue

    # 제목
    if s.startswith("### "):
        h = doc.add_heading(level=3); r = h.add_run(s[4:]); set_kr(r); r.font.color.rgb = RGBColor(0x3A,0x2A,0x20)
    elif s.startswith("## "):
        h = doc.add_heading(level=2); r = h.add_run(s[3:]); set_kr(r); r.font.color.rgb = ACCENT
    elif s.startswith("# "):
        h = doc.add_heading(level=1); r = h.add_run(s[2:]); set_kr(r); r.font.color.rgb = ACCENT
    elif s.startswith("- "):
        p = doc.add_paragraph(style="List Bullet"); add_runs(p, s[2:])
    elif s == "---" or s == "":
        if s == "":
            pass
    else:
        p = doc.add_paragraph(); add_runs(p, s)
    i += 1

out = HERE / "기획서_v2.docx"
doc.save(str(out))
print("saved", out)
