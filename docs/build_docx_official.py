# -*- coding: utf-8 -*-
"""공모전 공식 양식(원본 기획서 PDF) 형식 그대로 docx 생성 — 별도 디자인 없음."""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = Path(__file__).resolve().parent
IMG = HERE / "images"
FONT = "맑은 고딕"

doc = Document()
for sec in doc.sections:
    sec.top_margin = sec.bottom_margin = Inches(0.8)
    sec.left_margin = sec.right_margin = Inches(0.9)

st = doc.styles["Normal"]
st.font.name = FONT
st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
st.paragraph_format.space_after = Pt(4)
st.paragraph_format.line_spacing = 1.3


def kr(run):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def shade(el, fill):
    pr = el.get_or_add_tcPr() if el.tag.endswith("tc") else el.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    pr.append(shd)


def section_bar(num_title):
    """1) ... 형태 회색 바 섹션 제목."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
    shade(p._p, "D9D9D9")
    r = p.add_run(num_title); r.bold = True; r.font.size = Pt(11); kr(r)
    return p


def body(text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.4
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True; kr(r)
    r = p.add_run(text); kr(r)
    return p


def pic(fname, caption):
    if (IMG / fname).exists():
        doc.add_picture(str(IMG / fname), width=Inches(5.4))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption); r.font.size = Pt(9); kr(r)


def set_cell(cell, text, bold=False, fill=None, size=10):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); kr(r)
    if fill:
        shade(cell._tc, fill)


# ── 제목 (박스) ─────────────────────────────────────────────
title_tbl = doc.add_table(rows=1, cols=1); title_tbl.style = "Table Grid"
tc = title_tbl.cell(0, 0)
p = tc.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("「제4회 문화체육관광 인공지능·데이터 활용 공모전」기획서"); r.bold = True; r.font.size = Pt(15); kr(r)
p2 = tc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p2.add_run("- 문화데이터 활용 분야 -"); r.bold = True; r.font.size = Pt(12); kr(r)
doc.add_paragraph()

# ── 상단 정보 표 ────────────────────────────────────────────
info = doc.add_table(rows=3, cols=2); info.style = "Table Grid"
info.columns[0].width = Inches(1.4); info.columns[1].width = Inches(5.2)
set_cell(info.cell(0, 0), "공모 부문", bold=True, fill="F2F2F2")
set_cell(info.cell(0, 1), "제품·서비스 부문")
set_cell(info.cell(1, 0), "사 례 명", bold=True, fill="F2F2F2")
set_cell(info.cell(1, 1), "국가별 음식 기반 유사 한식 추천 및 관광 연계 데이터 매핑 (K-Food Match)")
set_cell(info.cell(2, 0), "제품·서비스 유형", bold=True, fill="F2F2F2")
typ = info.cell(2, 1)
lines = [
    "□ 모바일 앱(APP)   URL 주소 :",
    "☑ 웹   URL 주소 : https://sungchul02.github.io/KFood/",
    "□ 기타",
]
for i, t in enumerate(lines):
    pp = typ.paragraphs[0] if i == 0 else typ.add_paragraph()
    rr = pp.add_run(t); rr.font.size = Pt(10); kr(rr)
doc.add_paragraph()

# ── 1) 소개 및 목적 ─────────────────────────────────────────
section_bar("1) 제품·서비스 소개 및 목적(요약)")
body("최근 한국을 방문하는 외국인 관광객이 증가하면서 음식은 관광 만족도를 결정하는 중요한 요소로 자리잡고 있다. "
     "관광객들은 명소 방문뿐 아니라 현지 음식 체험을 여행의 주요 목적으로 인식하며, 음식 경험은 한국 문화에 대한 이해와 "
     "재방문 의사에도 큰 영향을 미친다. 그러나 한국 음식은 음식명만으로 재료와 맛을 직관적으로 파악하기 어렵다. "
     "'김치찌개', '육개장', '순두부찌개'는 한국인에게 익숙하지만, 외국인 관광객은 해당 음식이 맵거나 특정 향신료를 쓰는지, "
     "자신이 즐겨 먹는 음식과 유사한지 판단하기 어렵다. 이러한 정보 부족은 음식 선택을 부담스럽게 만들고, 결과적으로 "
     "소수의 유명 음식과 관광지에 소비가 집중되는 한계로 이어진다.")
body("본 서비스는 이 문제를 두 축으로 해결한다. 첫째, 관광객이 익숙한 자국 음식을 기준으로 재료와 맛이 유사한 한국 음식을 "
     "추천하여 한식에 대한 심리적 진입장벽을 낮춘다. 둘째, 추천된 한식을 실제로 즐길 수 있는 맛집과 관련 특산물·전통시장을 "
     "공공데이터로 안내하여, 음식 경험을 지역 관광 동선으로 확장한다.")
body("이를 위해 공공데이터(한식 레시피·해외 음식·관광정보)를 수집·가공하여 음식 간 관계를 구조화한 음식 데이터 허브를 "
     "구축하였으며, 본 기획서 제출 시점에 실제 동작하는 웹 서비스로 구현을 완료하였다.")

# ── 2) 상세 설명 ───────────────────────────────────────────
section_bar("2) 제품·서비스 상세 설명")
body("서비스 이용 흐름은 ① 세계지도(홈)에서 국가 선택 → ② 국가별 대표 음식 선택 → ③ 유사 한국 음식 추천 → "
     "④ 맛집·특산물 안내의 4단계로 구성된다.")
body("접속 시 드래그로 회전하는 세계 지도(지구본)가 나타나며, 서비스가 데이터를 보유한 37개국이 강조 표시된다. "
     "사용자가 자신의 국가를 선택하면 해당 국가의 대표 음식이 사진과 함께 제시되고, 평소 즐겨 먹는 음식을 고를 수 있다.",
     bold_lead="[화면 ① 홈 / ② 음식 선택]  ")
pic("01_globe.png", "[그림 1] 홈 화면 — 드래그로 회전하는 세계 지도에서 국가 선택")
pic("02_foods.png", "[그림 2] 선택한 국가의 대표 음식 선택 화면")
body("사용자가 음식을 선택하면 재료와 맛이 가장 비슷한 한국 음식을 유사도 순으로 추천한다. 각 추천 결과에는 일치율(%), "
     "맛 특성(매운맛·감칠맛·고소한맛·단맛 등), 그리고 선택한 음식과 한식이 공유하는 공통 재료가 함께 표시되어 추천 근거를 "
     "투명하게 제공한다. 이를 통해 사용자는 익숙한 음식과 왜 그 한식이 비슷한지 납득하며 새로운 음식을 시도할 수 있다.",
     bold_lead="[화면 ③ 유사 한식 추천]  ")
pic("03_recommend.png", "[그림 3] 유사 한식 추천 결과 — 일치율·맛 특성·공통 재료 제시")
body("각 추천 한식에서 '맛집·특산물 찾기'를 선택하면, 그 음식을 즐길 수 있는 음식점과 관련 특산물·전통시장 정보가 "
     "한국관광공사 관광정보(TourAPI)를 통해 사진·주소와 함께 제공되며, 네이버·카카오 지도로 바로 연결된다. 이로써 음식 "
     "추천이 실제 지역 방문과 소비로 이어지도록 설계하였다.",
     bold_lead="[화면 ④ 맛집·특산물 안내]  ")
pic("04_places.png", "[그림 4] 추천 한식의 맛집·특산물 안내 (한국관광공사 TourAPI)")

# ── 3) 차별성 ──────────────────────────────────────────────
section_bar("3) 제품·서비스 차별성")
body("기존 음식 추천 서비스는 주로 사용자 리뷰, 평점, 인기 순위를 기반으로 추천한다. 이러한 방식은 현지 사용자에게는 "
     "유용하나, 한국 음식에 대한 경험이 부족한 외국인 관광객에게는 실질적인 도움을 주기 어렵다.")
body("본 서비스의 가장 큰 차별점은, 음식과 음식 사이의 관계를 재료·맛 속성으로 구조화한 음식 데이터 허브를 직접 구축하고, "
     "언어가 다른 해외 음식과 한국 음식을 하나의 표준 재료 체계로 연결한다는 점이다. 예를 들어 영어 재료 'Soy Sauce'와 "
     "한글 재료 '간장'을 동일한 표준 토큰으로 매칭하여, 국가 간 음식 비교와 추천을 가능하게 한다.")
body("또한 단순 음식 추천에 그치지 않고, 추천 결과를 맛집·특산물 등 관광 정보와 연계한다는 점에서 차별화된다. 즉 본 "
     "서비스는 단순 추천 시스템이 아니라 식문화와 관광을 연결하는 데이터 플랫폼으로, 향후 한식 콘텐츠 제작, 관광 산업, "
     "식품 산업, 문화 연구 등 다양한 분야로 확장 가능한 높은 확장성을 가진다.")

# ── 4) 성과 및 기대효과 ─────────────────────────────────────
section_bar("4) 제품·서비스 성과 및 기대효과")
body("외국인 관광객은 자신의 식문화와 유사한 한국 음식을 추천받음으로써 낯선 음식에 대한 부담을 줄이고, 보다 적극적으로 "
     "한국 음식 문화를 경험할 수 있다. 이는 관광 만족도 향상으로 이어진다.", bold_lead="① 관광 활성화  ")
body("나아가 추천 한식의 맛집과 특산물 판매처를 전국 단위로 안내함으로써, 특정 유명 관광지에 집중된 소비를 다양한 지역으로 "
     "분산시키고 지역 상권과 전통시장 활성화에 기여할 수 있다. 음식 추천과 위치 정보를 결합하여 먹거리 중심의 여행 동선을 "
     "손쉽게 계획하도록 지원한다.")
body("음식의 맛·향·식감과 같은 비정형 정보를 데이터로 변환함으로써, 외국인이 한식을 자국 식문화의 언어로 이해하도록 돕는다. "
     "재료·맛으로 연결된 데이터 허브는 어느 나라 음식과 어떤 한식이 닮았는지를 보여주는 식문화 교류·연구의 기초 데이터가 "
     "된다.", bold_lead="② 식문화 진흥(한식 세계화)  ")
body("축적된 음식 관계 데이터는 한식 콘텐츠 제작, 식품 산업, 관광 마케팅 등으로 확장 가능하며, 데이터가 추가될수록 추천 "
     "정확도와 적용 국가가 함께 확대되어 한식 세계화의 기초 자료로 활용될 수 있다.")

# ── 5) 문화데이터 활용 ─────────────────────────────────────
section_bar("5) 문화데이터 활용")
dt = doc.add_table(rows=2, cols=4); dt.style = "Table Grid"
a = dt.cell(0, 0); a.merge(dt.cell(1, 0)); set_cell(a, "활용 데이터(명)", bold=True, fill="F2F2F2")
b = dt.cell(0, 1); b.merge(dt.cell(1, 1)); set_cell(b, "제공기관(명)", bold=True, fill="F2F2F2")
c = dt.cell(0, 2); c.merge(dt.cell(0, 3)); set_cell(c, "출처", bold=True, fill="F2F2F2")
set_cell(dt.cell(1, 2), "플랫폼(명)", bold=True, fill="F2F2F2")
set_cell(dt.cell(1, 3), "URL", bold=True, fill="F2F2F2")
rows = [
    ("조리식품의 레시피DB", "식품의약품안전처", "식품안전나라 / 공공데이터포털",
     "https://www.foodsafetykorea.go.kr (COOKRCP01)"),
    ("THEMEALDB", "THEMEALDB", "THEMEALDB", "https://www.themealdb.com/api.php"),
    ("국문 관광정보 서비스(GW)", "한국관광공사", "공공데이터포털(data.go.kr)",
     "https://www.data.go.kr/data/15101578/openapi.do"),
]
for d in rows:
    rc = dt.add_row().cells
    for j, v in enumerate(d):
        set_cell(rc[j], v, size=9)
doc.add_paragraph()
body("수집·가공 과정은 다음과 같다. 첫째, 식품안전나라 OpenAPI로 한식 레시피(유효 1,051건)를, THEMEALDB에서 해외 음식"
     "(671건, 37개국)을 수집하였다. 둘째, 레시피 원문에서 용량·구획어를 제거해 재료명만 추출하였다(예: '양파(50g)' → "
     "'양파'). 셋째, 핵심 재료를 표준 토큰으로 통일하여(예: '간장'='Soy Sauce'='soy_sauce') 언어가 다른 두 데이터를 "
     "연결하였다. 넷째, 재료로부터 매운맛·감칠맛·고소함 등 맛 속성을 산출하였다. 다섯째, 추천된 한식명에서 대표 요리 종류 "
     "키워드를 추출하여 한국관광공사 TourAPI로 맛집·특산물 정보를 실시간 조회한다. 문화데이터는 본 서비스에서 추천의 "
     "근거이자 관광 연계의 실데이터로서 핵심 역할을 수행한다.")

# ── 6) AI 기술 활용 ───────────────────────────────────────
section_bar("6) AI 기술 활용")
body("본 서비스는 음식의 맛·향·식감과 같은 비정형 특성을 정형 데이터로 변환하여 음식 간 비교와 추천을 가능하게 한다. "
     "표준 재료 토큰에 대한 가중 유사도(양념·맛 핵심 재료에 가중치 부여)와 맛 프로파일 유사도를 결합하여 추천 점수를 "
     "산출하며, 새로운 음식이 추가되면 동일한 추출·표준화 과정을 반복해 데이터 허브가 자동으로 확장되도록 설계하였다.")
body("현재는 규칙 기반으로 재료에서 맛 속성을 추출하고 있으며, 향후 생성형 AI(OpenAI API 등)를 활용한 속성 추출로 교체하여 "
     "향·식감·조리 특성까지 더 정밀하게 정형화할 계획이다. 이를 통해 사용자의 식문화적 배경을 고려한 개인화된 음식 추천과 "
     "관광 안내를 지속적으로 고도화할 수 있다.")

# ── 유의사항 ───────────────────────────────────────────────
doc.add_paragraph()
note = doc.add_table(rows=1, cols=1); note.style = "Table Grid"
nc = note.cell(0, 0)
items = [
    "※ 기획서 작성 시 유의사항",
    "  - 10장 내외 자유형식으로 작성하되, 1~6번 필수 항목은 반드시 명시하여 작성",
    "  - 제품·서비스 개발의 경우 모바일 앱/웹 등 실행 여부 확인 가능한 자료 첨부(캡처 이미지, URL 등)",
    "  - 이미지 파일은 문서 내 포함",
]
for i, t in enumerate(items):
    pp = nc.paragraphs[0] if i == 0 else nc.add_paragraph()
    rr = pp.add_run(t); rr.font.size = Pt(9); rr.bold = (i == 0); kr(rr)

out = HERE / "기획서_공식양식.docx"
doc.save(str(out))
print("saved", out)
